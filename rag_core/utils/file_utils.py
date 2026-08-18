"""相互独立的 TXT、PDF、DOCX 和 DOC 文件解析函数。"""

from pathlib import Path
import shutil
import subprocess
import tempfile

from rag_core.errors import DocumentParseError


def _existing_file(path: str | Path) -> Path:
    file_path = Path(path).expanduser().resolve()
    if not file_path.is_file():
        raise FileNotFoundError(f"文档不存在：{file_path}")
    return file_path


def _ensure_text(text: str, file_path: Path) -> str:
    if not text or not text.strip():
        raise DocumentParseError(f"文档没有可读取的文本：{file_path}")
    return text


def read_txt(path: str | Path) -> str:
    """读取 TXT，依次尝试 UTF-8-SIG、GB18030 和 UTF-16。"""
    file_path = _existing_file(path)
    errors: list[str] = []
    for encoding in ("utf-8-sig", "gb18030", "utf-16"):
        try:
            return _ensure_text(file_path.read_text(encoding=encoding), file_path)
        except UnicodeError as exc:
            errors.append(f"{encoding}: {exc}")
    raise DocumentParseError(
        f"无法识别 TXT 编码：{file_path}；已尝试 UTF-8、GB18030、UTF-16"
    )


def read_pdf(path: str | Path) -> str:
    """使用 pypdf 提取 PDF 中可复制的文本。"""
    file_path = _existing_file(path)
    try:
        from pypdf import PdfReader

        reader = PdfReader(str(file_path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return _ensure_text("\n\n".join(pages), file_path)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"PDF 解析失败：{file_path}：{exc}") from exc


def read_docx(path: str | Path) -> str:
    """使用 python-docx 读取 DOCX 段落和表格文本。"""
    file_path = _existing_file(path)
    try:
        from docx import Document

        document = Document(str(file_path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return _ensure_text("\n".join(parts), file_path)
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"DOCX 解析失败：{file_path}：{exc}") from exc


def read_doc(path: str | Path) -> str:
    """通过 LibreOffice 将旧版 DOC 临时转换为 DOCX 后读取。"""
    file_path = _existing_file(path)
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise DocumentParseError(
            "解析旧版 .doc 需要安装 LibreOffice，并确保 soffice 命令位于 PATH 中"
        )
    with tempfile.TemporaryDirectory(prefix="rag_doc_") as temp_dir:
        result = subprocess.run(
            [
                soffice,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                temp_dir,
                str(file_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,
            check=False,
        )
        converted = Path(temp_dir) / f"{file_path.stem}.docx"
        if result.returncode != 0 or not converted.exists():
            detail = result.stderr.strip() or result.stdout.strip() or "未知错误"
            raise DocumentParseError(f"DOC 转换失败：{file_path}：{detail}")
        return read_docx(converted)


def read_document(path: str | Path) -> str:
    """按扩展名分派到对应的原子解析函数。"""
    file_path = _existing_file(path)
    readers = {
        ".txt": read_txt,
        ".pdf": read_pdf,
        ".docx": read_docx,
        ".doc": read_doc,
    }
    reader = readers.get(file_path.suffix.lower())
    if not reader:
        supported = ", ".join(sorted(readers))
        raise DocumentParseError(
            f"不支持的文档格式：{file_path.suffix or '<无扩展名>'}；支持 {supported}"
        )
    return reader(file_path)
